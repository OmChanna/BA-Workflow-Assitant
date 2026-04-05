# ══════════════════════════════════════════════════════════════════════════════
#  INGESTION PIPELINE — Parse, Chunk, Embed, Store
#  Handles document uploads for both structure_examples and domain_knowledge
#  Used by: pages/admin.py
# ══════════════════════════════════════════════════════════════════════════════

from __future__ import annotations
import io
import json
from datetime import datetime
from typing import Optional

from knowledge_store import (
    COLLECTION_STRUCTURE, COLLECTION_DOMAIN,
    chunk_text, store_chunks, generate_doc_id,
    delete_document, ensure_collections,
)


# ── Document Parsers (reuse backend parsers + extend) ─────────────────────────

def parse_uploaded_file(uploaded_file) -> dict:
    """Parse an uploaded file into text. Returns {filename, text, file_type, success, error, page_count}.
    Accepts Streamlit UploadedFile objects.
    """
    filename = uploaded_file.name
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)

    result = {
        "filename": filename, "text": "", "file_type": "",
        "success": False, "error": None, "page_count": 0,
    }

    try:
        if filename.lower().endswith(".txt") or filename.lower().endswith(".md"):
            result["text"] = file_bytes.decode("utf-8", errors="replace")
            result["file_type"] = "TXT"
            result["success"] = True
            result["page_count"] = max(1, len(result["text"]) // 3000)

        elif filename.lower().endswith(".pdf"):
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                pages = []
                for page in reader.pages:
                    pages.append(page.extract_text() or "")
                result["text"] = "\n\n".join(pages)
                result["file_type"] = "PDF"
                result["success"] = True
                result["page_count"] = len(reader.pages)
            except ImportError:
                result["error"] = "PyPDF2 not installed."
            except Exception as e:
                result["error"] = f"PDF parsing failed: {str(e)}"

        elif filename.lower().endswith(".docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            paragraphs.append(row_text)
                result["text"] = "\n".join(paragraphs)
                result["file_type"] = "Word"
                result["success"] = True
                result["page_count"] = max(1, len(result["text"]) // 3000)
            except ImportError:
                result["error"] = "python-docx not installed."
            except Exception as e:
                result["error"] = f"Word parsing failed: {str(e)}"

        elif filename.lower().endswith((".xlsx", ".xls")):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
                sheets_text = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        row_vals = [str(c) if c is not None else "" for c in row]
                        if any(v.strip() for v in row_vals):
                            rows.append(" | ".join(row_vals))
                    if rows:
                        sheets_text.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
                result["text"] = "\n\n".join(sheets_text)
                result["file_type"] = "Excel"
                result["success"] = True
                result["page_count"] = len(wb.sheetnames)
            except ImportError:
                result["error"] = "openpyxl not installed."
            except Exception as e:
                result["error"] = f"Excel parsing failed: {str(e)}"

        elif filename.lower().endswith(".csv"):
            result["text"] = file_bytes.decode("utf-8", errors="replace")
            result["file_type"] = "CSV"
            result["success"] = True
            result["page_count"] = 1

        elif filename.lower().endswith(".json"):
            text = file_bytes.decode("utf-8", errors="replace")
            try:
                data = json.loads(text)
                result["text"] = json.dumps(data, indent=2)
            except json.JSONDecodeError:
                result["text"] = text
            result["file_type"] = "JSON"
            result["success"] = True
            result["page_count"] = 1

        else:
            try:
                result["text"] = file_bytes.decode("utf-8", errors="replace")
                result["file_type"] = "Text"
                result["success"] = True
                result["page_count"] = 1
            except Exception:
                result["error"] = f"Unsupported file type: {filename}"

    except Exception as e:
        result["error"] = f"File processing failed: {str(e)}"

    return result


# ── Agent ID Mapping (for structure examples) ─────────────────────────────────

AGENT_DISPLAY_MAP = {
    "A01": {"name": "Stakeholder Analysis", "tab": "Pre-Discovery"},
    "A04": {"name": "BRD", "tab": "Requirements"},
    "A15": {"name": "FRD", "tab": "Requirements"},
    "A05": {"name": "NFR Library", "tab": "Requirements"},
    "A09_ASIS": {"name": "AS-IS Process Flow", "tab": "Requirements"},
    "A09_TOBE": {"name": "TO-BE Process Flow", "tab": "Requirements"},
    "A09_FSD": {"name": "Functional Specification", "tab": "Analysis"},
    "A08": {"name": "Risk & Priority", "tab": "Prioritisation"},
    "A10": {"name": "Change Impact", "tab": "Analysis"},
    "A06": {"name": "Traceability Matrix", "tab": "Analysis"},
    "A14": {"name": "Agile/Sprint", "tab": "Agile/Scrum"},
    "A11": {"name": "Test Scripts", "tab": "Testing"},
    "A13": {"name": "Handover Docs", "tab": "Handover"},
}

# Domain → Subdomain tree
DOMAIN_TREE = {
    "life_sciences": {
        "name": "Life Sciences",
        "subdomains": {
            "commercial": "Commercial (CRM, HCP, Sales Force)",
            "regulatory": "Regulatory Affairs",
            "manufacturing": "Manufacturing & Supply",
            "clinical": "Clinical Operations",
            "pharmacovigilance": "Pharmacovigilance & Safety",
            "medical_affairs": "Medical Affairs",
        },
    },
    "healthcare": {
        "name": "Healthcare",
        "subdomains": {
            "ehr": "EHR/EMR Systems",
            "revenue_cycle": "Revenue Cycle Management",
            "patient_engagement": "Patient Engagement",
            "population_health": "Population Health",
        },
    },
    "financial_services": {
        "name": "Financial Services",
        "subdomains": {
            "banking": "Banking",
            "insurance": "Insurance",
            "capital_markets": "Capital Markets",
            "compliance": "Financial Compliance",
        },
    },
    "general_it": {
        "name": "General IT",
        "subdomains": {
            "crm": "CRM Implementations",
            "erp": "ERP Systems",
            "data_analytics": "Data & Analytics",
            "cloud_migration": "Cloud Migration",
        },
    },
}


# ── Ingestion Functions ───────────────────────────────────────────────────────

def ingest_structure_example(
    api_key: str,
    uploaded_file,
    agent_id: str,
    domain: str = "life_sciences",
    subdomain: str = "commercial",
) -> dict:
    """Ingest a structure/template example document for a specific agent.

    Pipeline: parse → chunk → embed → store in structure_examples collection.

    Returns:
        {success, doc_id, filename, chunk_count, error}
    """
    # Step 1: Parse
    parsed = parse_uploaded_file(uploaded_file)
    if not parsed["success"]:
        return {"success": False, "error": parsed["error"], "doc_id": None,
                "filename": parsed["filename"], "chunk_count": 0}

    if len(parsed["text"].strip()) < 50:
        return {"success": False, "error": "Document too short (< 50 chars).",
                "doc_id": None, "filename": parsed["filename"], "chunk_count": 0}

    # Step 2: Generate doc ID
    doc_id = generate_doc_id(
        filename=parsed["filename"],
        collection=COLLECTION_STRUCTURE,
        agent_id=agent_id,
        domain=domain,
        subdomain=subdomain,
    )

    # Step 3: Delete existing chunks for this doc (re-upload = replace)
    try:
        delete_document(COLLECTION_STRUCTURE, doc_id)
    except Exception:
        pass  # Collection might not exist yet

    # Step 4: Chunk
    chunks = chunk_text(parsed["text"])
    if not chunks:
        return {"success": False, "error": "No chunks generated from document.",
                "doc_id": doc_id, "filename": parsed["filename"], "chunk_count": 0}

    # Step 5: Metadata
    metadata = {
        "doc_id": doc_id,
        "filename": parsed["filename"],
        "file_type": parsed["file_type"],
        "agent_id": agent_id,
        "domain": domain,
        "subdomain": subdomain,
        "page_count": parsed["page_count"],
        "total_chunks": len(chunks),
        "knowledge_type": "structure_example",
    }

    # Step 6: Embed + Store
    try:
        ensure_collections()
        stored = store_chunks(api_key, COLLECTION_STRUCTURE, chunks, metadata)
        return {
            "success": True, "doc_id": doc_id, "filename": parsed["filename"],
            "chunk_count": stored, "error": None,
        }
    except Exception as e:
        return {
            "success": False, "doc_id": doc_id, "filename": parsed["filename"],
            "chunk_count": 0, "error": f"Storage failed: {str(e)}",
        }


def ingest_domain_knowledge(
    api_key: str,
    uploaded_file,
    domain: str,
    subdomain: str,
) -> dict:
    """Ingest a domain knowledge document (SOP, regulation, style guide, etc.).

    Pipeline: parse → chunk → embed → store in domain_knowledge collection.

    Returns:
        {success, doc_id, filename, chunk_count, error}
    """
    # Step 1: Parse
    parsed = parse_uploaded_file(uploaded_file)
    if not parsed["success"]:
        return {"success": False, "error": parsed["error"], "doc_id": None,
                "filename": parsed["filename"], "chunk_count": 0}

    if len(parsed["text"].strip()) < 50:
        return {"success": False, "error": "Document too short (< 50 chars).",
                "doc_id": None, "filename": parsed["filename"], "chunk_count": 0}

    # Step 2: Generate doc ID
    doc_id = generate_doc_id(
        filename=parsed["filename"],
        collection=COLLECTION_DOMAIN,
        domain=domain,
        subdomain=subdomain,
    )

    # Step 3: Delete existing (re-upload = replace)
    try:
        delete_document(COLLECTION_DOMAIN, doc_id)
    except Exception:
        pass

    # Step 4: Chunk
    chunks = chunk_text(parsed["text"])
    if not chunks:
        return {"success": False, "error": "No chunks generated from document.",
                "doc_id": doc_id, "filename": parsed["filename"], "chunk_count": 0}

    # Step 5: Metadata
    metadata = {
        "doc_id": doc_id,
        "filename": parsed["filename"],
        "file_type": parsed["file_type"],
        "agent_id": "",  # domain knowledge is not agent-specific
        "domain": domain,
        "subdomain": subdomain,
        "page_count": parsed["page_count"],
        "total_chunks": len(chunks),
        "knowledge_type": "domain_knowledge",
    }

    # Step 6: Embed + Store
    try:
        ensure_collections()
        stored = store_chunks(api_key, COLLECTION_DOMAIN, chunks, metadata)
        return {
            "success": True, "doc_id": doc_id, "filename": parsed["filename"],
            "chunk_count": stored, "error": None,
        }
    except Exception as e:
        return {
            "success": False, "doc_id": doc_id, "filename": parsed["filename"],
            "chunk_count": 0, "error": f"Storage failed: {str(e)}",
        }
